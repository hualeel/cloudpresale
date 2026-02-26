"""
Word 技术方案生成器 (python-docx)
输入：Solution.content（6个Agent输出汇总）+ 需求信息
输出：.docx 字节流
"""
import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 颜色常量 ──────────────────────────────────────
COLOR_PRIMARY = RGBColor(0x0A, 0x84, 0xFF)   # 科技蓝
COLOR_DARK    = RGBColor(0x1A, 0x1A, 0x2E)   # 深色标题
COLOR_GRAY    = RGBColor(0x60, 0x60, 0x70)   # 灰色辅助


def _set_paragraph_spacing(para, before=0, after=6, line_rule=None):
    """设置段落间距。"""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _add_heading(doc: Document, text: str, level: int = 1):
    """添加标题段落，带自定义颜色。"""
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = COLOR_DARK if level == 1 else COLOR_PRIMARY
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)
    _set_paragraph_spacing(para, before=12, after=6)
    return para


def _add_body(doc: Document, text: str, bold_start: bool = False):
    """添加正文段落，支持 **粗体** Markdown 格式。"""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, before=0, after=4)
    # 解析 **bold** 语法
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        run.font.size = Pt(10.5)
    return para


def _parse_markdown_section(doc: Document, content: str):
    """
    将单个 Agent 的 Markdown 内容解析并写入 Word。
    支持: ## 标题, - 列表, | 表格, 正文
    """
    if not content:
        return

    lines = content.strip().split('\n')
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        # 提取数据行（跳过分隔行 ---）
        data_rows = [l for l in table_lines if not re.match(r'^\|[-:\s|]+\|$', l)]
        if len(data_rows) < 2:
            table_lines = []
            return
        headers = [c.strip() for c in data_rows[0].strip('|').split('|')]
        body_rows = []
        for row_line in data_rows[1:]:
            cells = [c.strip() for c in row_line.strip('|').split('|')]
            body_rows.append(cells)

        table = doc.add_table(rows=1 + len(body_rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        # 表头
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            if i < len(hdr_row.cells):
                cell = hdr_row.cells[i]
                cell.text = h
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = COLOR_PRIMARY
        # 数据行
        for ri, row_data in enumerate(body_rows):
            tr = table.rows[ri + 1]
            for ci, cell_val in enumerate(row_data):
                if ci < len(tr.cells):
                    tr.cells[ci].text = cell_val
        doc.add_paragraph()
        table_lines = []

    for line in lines:
        # 检测表格行
        if line.strip().startswith('|'):
            table_lines.append(line)
            continue
        else:
            flush_table()

        if line.startswith('## '):
            _add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            _add_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            _add_heading(doc, line[5:], level=4)
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            _add_body_to_para(para, line[2:])
            _set_paragraph_spacing(para, before=0, after=2)
        elif line.strip() == '':
            doc.add_paragraph()
        else:
            _add_body(doc, line)

    flush_table()


def _add_body_to_para(para, text: str):
    """向已有段落添加 run（支持 **bold**）。"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def _add_cover_page(doc: Document, title: str, customer: str, version: str, date_str: str):
    """添加封面页。"""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK
    _set_paragraph_spacing(p_title, before=0, after=12)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_sub.add_run(f"客户：{customer}")
    run2.font.size = Pt(14)
    run2.font.color.rgb = COLOR_GRAY

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p_ver.add_run(f"版本 {version}  |  {date_str}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = COLOR_GRAY

    doc.add_page_break()


def generate_word_tech(
    solution_content: dict,
    requirement: dict,
    customer_name: str,
    version: str,
) -> bytes:
    """
    生成 Word 技术方案文档。

    Args:
        solution_content: Solution.content（各 Agent 输出）
        requirement: Requirement 信息字典（title, content）
        customer_name: 客户名称
        version: 版本号（如 "1.0"）

    Returns:
        .docx 文件字节流
    """
    doc = Document()

    # 页面设置（A4）
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)

    date_str = datetime.now().strftime("%Y年%m月%d日")
    doc_title = requirement.get("title", "云原生技术方案")

    # 封面
    _add_cover_page(doc, doc_title, customer_name, version, date_str)

    # 文档说明
    _add_heading(doc, "文档说明", level=1)
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ("文档标题", doc_title),
        ("客户名称", customer_name),
        ("文档版本", f"v{version}"),
        ("编制日期", date_str),
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()

    # 章节顺序
    SECTIONS = [
        ("arch",      "1. 架构设计方案"),
        ("sizing",    "2. 资源规格设计"),
        ("security",  "3. 安全合规方案"),
        ("migration", "4. 应用迁移路径"),
        ("plan",      "5. 项目实施计划"),
        ("pricing",   "6. 商务报价方案"),
    ]

    for key, heading in SECTIONS:
        agent_result = solution_content.get(key)
        if not agent_result:
            continue
        content_text = agent_result.get("content", "") if isinstance(agent_result, dict) else str(agent_result)
        _add_heading(doc, heading, level=1)
        _parse_markdown_section(doc, content_text)
        doc.add_page_break()

    # 保存到字节流
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
